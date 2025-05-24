
from docx import Document
from datetime import datetime
import io

def generate_docx(name, age, gender, symptoms, activity, exposure, smoking, hrv):
    doc = Document()
    doc.add_heading('AMAL Radiology Report', 0)
    doc.add_paragraph(f'Date: {datetime.today().strftime("%Y-%m-%d")}')
    doc.add_paragraph(f'Patient Name: {name}')
    doc.add_paragraph(f'Age: {age}   Gender: {gender}')
    doc.add_heading('Clinical Information', level=1)
    doc.add_paragraph(f'Symptoms: {symptoms or "N/A"}')
    doc.add_paragraph(f'Physical Activity: {activity}')
    doc.add_paragraph(f'Exposure: {exposure or "N/A"}')
    doc.add_paragraph(f'Smoking History: {smoking}')
    doc.add_paragraph(f'Heart Rate Variability: {hrv or "N/A"} ms')
    doc.add_heading('AI Diagnosis Summary', level=1)
    doc.add_paragraph('Chest X‑ray indicates areas of consolidation consistent with Pneumonia.')
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph('- Correlate clinically.')
    doc.add_paragraph('- Consider antibiotic therapy and treatment.')
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
